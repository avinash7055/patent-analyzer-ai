import zipfile
import xml.etree.ElementTree as ET
from pathlib import Path
from typing import Optional
from dataclasses import dataclass, field

try:
    from docx import Document as DocxDocument
    HAS_PYTHON_DOCX = True
except ImportError:
    HAS_PYTHON_DOCX = False


WML_NS = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"


@dataclass
class Paragraph:
    index: int
    style: str
    text: str


@dataclass
class Table:
    index: int
    rows: list[list[str]]

    @property
    def num_rows(self) -> int:
        return len(self.rows)

    @property
    def num_cols(self) -> int:
        return len(self.rows[0]) if self.rows else 0


@dataclass
class ProcessedDocument:
    filename: str
    paragraphs: list[Paragraph] = field(default_factory=list)
    tables: list[Table] = field(default_factory=list)
    full_text: str = ""
    sections: dict[str, str] = field(default_factory=dict)
    doc_type: str = ""


class DocumentProcessor:

    def process(self, filepath: str | Path) -> ProcessedDocument:
        filepath = Path(filepath)
        try:
            return self._process_with_python_docx(filepath)
        except Exception:
            return self._process_with_xml_fallback(filepath)

    def _process_with_python_docx(self, filepath: Path) -> ProcessedDocument:
        if not HAS_PYTHON_DOCX:
            raise RuntimeError("python-docx not available")

        doc = DocxDocument(str(filepath))
        paragraphs = []
        for i, para in enumerate(doc.paragraphs):
            if para.text.strip():
                paragraphs.append(Paragraph(
                    index=i,
                    style=para.style.name if para.style else "Normal",
                    text=para.text
                ))

        tables = []
        for i, table in enumerate(doc.tables):
            rows = []
            for row in table.rows:
                cells = [cell.text.strip() for cell in row.cells]
                rows.append(cells)
            tables.append(Table(index=i, rows=rows))

        full_text = "\n".join(p.text for p in paragraphs)
        sections = self._extract_sections(paragraphs)
        doc_type = self._detect_doc_type(full_text)

        return ProcessedDocument(
            filename=filepath.name,
            paragraphs=paragraphs,
            tables=tables,
            full_text=full_text,
            sections=sections,
            doc_type=doc_type
        )

    def _process_with_xml_fallback(self, filepath: Path) -> ProcessedDocument:
        zf = zipfile.ZipFile(str(filepath))
        xml_content = zf.read("word/document.xml")
        zf.close()

        paragraphs = self._extract_paragraphs_from_xml(xml_content)
        tables = self._extract_tables_from_xml(xml_content)
        full_text = "\n".join(p.text for p in paragraphs)
        sections = self._extract_sections(paragraphs)
        doc_type = self._detect_doc_type(full_text)

        return ProcessedDocument(
            filename=filepath.name,
            paragraphs=paragraphs,
            tables=tables,
            full_text=full_text,
            sections=sections,
            doc_type=doc_type
        )

    def _extract_paragraphs_from_xml(self, xml_content: bytes) -> list[Paragraph]:
        root = ET.fromstring(xml_content)
        paragraphs = []
        idx = 0

        for p_elem in root.iter(f"{{{WML_NS}}}p"):
            texts = []
            for t_elem in p_elem.iter(f"{{{WML_NS}}}t"):
                if t_elem.text:
                    texts.append(t_elem.text)
            para_text = "".join(texts)
            if para_text.strip():
                style = self._get_style_from_xml(p_elem)
                paragraphs.append(Paragraph(index=idx, style=style, text=para_text))
                idx += 1

        return paragraphs

    def _extract_tables_from_xml(self, xml_content: bytes) -> list[Table]:
        root = ET.fromstring(xml_content)
        tables = []
        idx = 0

        for tbl in root.iter(f"{{{WML_NS}}}tbl"):
            rows = []
            for tr in tbl.iter(f"{{{WML_NS}}}tr"):
                cells = []
                for tc in tr.iter(f"{{{WML_NS}}}tc"):
                    cell_texts = []
                    for t in tc.iter(f"{{{WML_NS}}}t"):
                        if t.text:
                            cell_texts.append(t.text)
                    cells.append(" ".join(cell_texts))
                rows.append(cells)
            tables.append(Table(index=idx, rows=rows))
            idx += 1

        return tables

    def _get_style_from_xml(self, p_elem) -> str:
        ppr = p_elem.find(f"{{{WML_NS}}}pPr")
        if ppr is not None:
            pstyle = ppr.find(f"{{{WML_NS}}}pStyle")
            if pstyle is not None:
                return pstyle.get(f"{{{WML_NS}}}val", "Normal")
        return "Normal"

    def _extract_sections(self, paragraphs: list[Paragraph]) -> dict[str, str]:
        sections = {}
        current_section = None
        current_text = []

        for p in paragraphs:
            if self._is_section_heading(p):
                if current_section:
                    sections[current_section] = "\n".join(current_text)
                current_section = p.text.strip()
                current_text = []
            elif current_section:
                current_text.append(p.text)

        if current_section:
            sections[current_section] = "\n".join(current_text)

        return sections

    def _is_section_heading(self, p: Paragraph) -> bool:
        if "heading" in p.style.lower():
            return True
        import re
        if re.match(r"^\d+\.\d*\s", p.text.strip()):
            return True
        return False

    def _detect_doc_type(self, text: str) -> str:
        text_lower = text.lower()
        if "invention disclosure" in text_lower:
            return "IDF"
        if "patentability" in text_lower or "prior art" in text_lower:
            return "PR"
        return "UNKNOWN"


def extract_prior_art_sections(doc: ProcessedDocument) -> dict[str, str]:
    prior_arts = {}
    current_ref = None
    current_text = []

    for p in doc.paragraphs:
        text = p.text.strip()
        import re
        ref_match = re.match(r"Ref\s*(\d+)\s*\(D\s*(\d+)\)", text, re.IGNORECASE)
        if ref_match:
            if current_ref:
                prior_arts[current_ref] = "\n".join(current_text)
            current_ref = f"D{ref_match.group(2)}"
            current_text = [text]
        elif current_ref:
            current_text.append(text)

    if current_ref:
        prior_arts[current_ref] = "\n".join(current_text)

    if not prior_arts:
        prior_arts = _extract_prior_art_from_tables(doc)

    return prior_arts


def _extract_prior_art_from_tables(doc: ProcessedDocument) -> dict[str, str]:
    prior_arts = {}
    for table in doc.tables:
        if table.num_rows < 3 or table.num_cols < 3:
            continue
        for row in table.rows:
            row_text = " ".join(row)
            import re
            ref_match = re.search(r"(D\d+)\.\s*([\w\d]+)", row_text)
            if ref_match:
                ref_id = ref_match.group(1)
                prior_arts[ref_id] = row_text
    return prior_arts


def extract_pr_elements_text(doc: ProcessedDocument) -> str:
    elements_text = []
    capture = False

    for p in doc.paragraphs:
        text = p.text.strip()
        if "technical elements" in text.lower() or "element a" in text.lower():
            capture = True
        if capture:
            elements_text.append(text)
        if "references provided" in text.lower() and capture:
            break

    if not elements_text:
        for p in doc.paragraphs:
            text = p.text.strip()
            if text.startswith("Element") or text.startswith("A1") or text.startswith("A2") or text.startswith("B1"):
                elements_text.append(text)

    return "\n".join(elements_text)
